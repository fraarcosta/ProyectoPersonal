"""
FastAPI entry point — agent-oferta.
Endpoints:
  GET  /health
  POST /detect-incoherencias  — detecta incoherencias entre PCAP/PPT (multi-fichero)
  POST /generate-index        — genera índice de la oferta técnica (multi-fichero)
"""
import io
import logging
from typing import List
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from shared.factory import get_llm, get_memory_store, get_settings
from src.agent import OfertaAgent

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

settings = get_settings()
_agent: OfertaAgent | None = None


@asynccontextmanager
async def lifespan(app: FastAPI):
    global _agent
    _agent = OfertaAgent(llm=get_llm(), memory=get_memory_store())
    yield


app = FastAPI(
    title="Akena — Agent Oferta",
    version=settings.agent_version,
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Models ──────────────────────────────────────────────────────────────────

class IncoherenciasResponse(BaseModel):
    items: list
    agent: str = "agent-oferta"

class IndexResponse(BaseModel):
    index: str
    agent: str = "agent-oferta"

# ─── Text extraction helpers ─────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(p for p in parts if p.strip())
    except ImportError:
        raise HTTPException(status_code=422, detail="Librería pypdf no instalada.")
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Error leyendo PDF: {exc}")


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_texts = []
        for table in doc.tables:
            rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
            table_texts.append("\n".join(rows))
        return "\n\n".join(paragraphs + table_texts)
    except ImportError:
        raise HTTPException(status_code=422, detail="Librería python-docx no instalada.")
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Error leyendo DOCX: {exc}")


async def _extract_files(files: List[UploadFile]) -> tuple[str, str]:
    """Extrae texto de varios ficheros y los concatena. Retorna (texto_combinado, nombres)."""
    parts: list[str] = []
    names: list[str] = []
    for upload in files:
        file_name = upload.filename or ""
        ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
        if ext not in ("pdf", "docx"):
            raise HTTPException(status_code=415, detail=f"'{file_name}' no admitido. Solo .pdf y .docx.")
        data = await upload.read()
        if len(data) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail=f"'{file_name}' supera 50 MB.")
        text = _extract_pdf(data) if ext == "pdf" else _extract_docx(data)
        if text.strip():
            parts.append(f"=== DOCUMENTO: {file_name} ===\n\n{text}")
            names.append(file_name)
    if not parts:
        raise HTTPException(status_code=422, detail="No se pudo extraer texto de ningún fichero.")
    combined = "\n\n" + ("\n\n" + "─" * 60 + "\n\n").join(parts)
    return combined, " + ".join(names)

# ─── Endpoints ────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "agent": settings.agent_name, "version": settings.agent_version}


@app.post("/detect-incoherencias", response_model=IncoherenciasResponse)
async def detect_incoherencias(
    files: List[UploadFile] = File(...),
    session_id: str = Form("default"),
):
    """Detecta incoherencias, contradicciones y ambigüedades en el pliego."""
    text, names = await _extract_files(files)
    logging.info("detect-incoherencias: %d chars, files=%s, session=%s", len(text), names, session_id)
    assert _agent is not None
    try:
        items = await _agent.detect_incoherencias(
            document_text=text,
            session_id=session_id,
            file_name=names,
        )
    except Exception as exc:
        logging.error("Error en /detect-incoherencias: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    return IncoherenciasResponse(items=items)


@app.post("/generate-index", response_model=IndexResponse)
async def generate_index(
    files: List[UploadFile] = File(...),
    session_id: str = Form("default"),
):
    """Genera el índice estructurado de la oferta técnica a partir del pliego."""
    text, names = await _extract_files(files)
    logging.info("generate-index: %d chars, files=%s, session=%s", len(text), names, session_id)
    assert _agent is not None
    try:
        index = await _agent.generate_index(
            document_text=text,
            session_id=session_id,
            file_name=names,
        )
    except Exception as exc:
        logging.error("Error en /generate-index: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    return IndexResponse(index=index)
