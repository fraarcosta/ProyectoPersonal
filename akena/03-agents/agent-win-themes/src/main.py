"""
FastAPI entry point — agent-win-themes.
Endpoints:
  GET  /health
  GET  /.well-known/agent-card.json
  POST /generate-win-themes  — genera Win Themes por sección del índice a partir del pliego
"""
import io
import logging
from typing import List
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from shared.factory import get_llm, get_memory_store, get_settings
from src.agent import WinThemesAgent

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

settings = get_settings()
_agent: WinThemesAgent | None = None


@asynccontextmanager
async def lifespan(app: FastAPI):
    global _agent
    _agent = WinThemesAgent(llm=get_llm(), memory=get_memory_store())
    logging.info("[agent-win-themes] Agente inicializado — modelo: %s", get_llm().model_id)
    yield


app = FastAPI(
    title="Akena — Agent Win Themes",
    version=settings.agent_version,
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─── Models ──────────────────────────────────────────────────────────────────

class WinThemesResponse(BaseModel):
    sections: dict[str, str]
    agent: str = "agent-win-themes"


# ─── Text extraction helpers ──────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(p for p in parts if p.strip())
    except ImportError:
        raise HTTPException(status_code=422, detail="Librería pypdf no instalada.")
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Error leyendo PDF: {exc}")


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_texts = []
        for table in doc.tables:
            rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
            table_texts.append("\n".join(rows))
        return "\n\n".join(paragraphs + table_texts)
    except ImportError:
        raise HTTPException(status_code=422, detail="Librería python-docx no instalada.")
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Error leyendo DOCX: {exc}")


async def _extract_files(files: List[UploadFile]) -> tuple[str, str]:
    """Extrae y concatena texto de varios ficheros. Retorna (texto_combinado, nombres)."""
    parts: list[str] = []
    names: list[str] = []
    for upload in files:
        file_name = upload.filename or ""
        ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
        if ext not in ("pdf", "docx"):
            raise HTTPException(
                status_code=415,
                detail=f"'{file_name}' no admitido. Solo .pdf y .docx.",
            )
        data = await upload.read()
        if len(data) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail=f"'{file_name}' supera 50 MB.")
        text = _extract_pdf(data) if ext == "pdf" else _extract_docx(data)
        if text.strip():
            parts.append(f"=== DOCUMENTO: {file_name} ===\n\n{text}")
            names.append(file_name)
    if not parts:
        raise HTTPException(status_code=422, detail="No se pudo extraer texto de ningún fichero.")
    combined = "\n\n" + ("\n\n" + "─" * 60 + "\n\n").join(parts)
    return combined, " + ".join(names)


# ─── Endpoints ────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "agent": settings.agent_name, "version": settings.agent_version}


@app.get("/.well-known/agent-card.json")
async def agent_card():
    return {
        "name": settings.agent_name,
        "version": settings.agent_version,
        "type": "specialized",
        "description": "Genera Win Themes por sección del índice de la oferta técnica para maximizar la puntuación técnica en licitaciones públicas españolas.",
        "skills": [
            {
                "id": "generate_win_themes",
                "name": "Win Themes Generator",
                "description": "Analiza el pliego (PCAP, PPT, Anexos) junto con el índice validado y genera argumentos diferenciadores por apartado para maximizar la puntuación técnica.",
            }
        ],
    }


@app.post("/generate-win-themes", response_model=WinThemesResponse)
async def generate_win_themes(
    files: List[UploadFile] = File(...),
    index: str = Form(...),
    session_id: str = Form("default"),
):
    """
    Genera Win Themes para cada sección L1 del índice a partir del pliego.

    - files: documentos del pliego (PCAP, PPT, Anexos) en PDF o DOCX
    - index: texto completo del índice validado de la oferta (numeración decimal)
    - session_id: ID de sesión para memoria conversacional
    """
    if not index.strip():
        raise HTTPException(status_code=422, detail="El índice no puede estar vacío.")

    doc_text, file_names = await _extract_files(files)
    logging.info(
        "generate-win-themes: %d chars pliego, %d chars índice, files=%s, session=%s",
        len(doc_text), len(index), file_names, session_id,
    )
    # Log primeras líneas del índice para diagnóstico de formato
    logging.info("ÍNDICE (primeras 600 chars):\n%s", index[:600])

    assert _agent is not None
    try:
        sections = await _agent.generate_win_themes(
            document_text=doc_text,
            index_text=index,
            session_id=session_id,
            file_names=file_names,
        )
    except ValueError as exc:
        logging.error("Error de parseo en /generate-win-themes: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    except Exception as exc:
        logging.error("Error en /generate-win-themes: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))

    return WinThemesResponse(sections=sections)
