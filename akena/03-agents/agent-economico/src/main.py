"""
FastAPI entry point — agent-economico.
Endpoints:
  GET  /health
  POST /extract-formula  — extrae fórmula económica del PCAP (multi-fichero)
  POST /simulate         — simula escenarios de precio con la fórmula extraída
"""
import io
import logging
from typing import List
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from shared.factory import get_llm, get_memory_store, get_settings
from src.agent import EconomicoAgent

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

settings = get_settings()
_agent: EconomicoAgent | None = None


@asynccontextmanager
async def lifespan(app: FastAPI):
    global _agent
    _agent = EconomicoAgent(llm=get_llm(), memory=get_memory_store())
    yield


app = FastAPI(
    title="Akena — Agent Económico",
    version=settings.agent_version,
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Models ──────────────────────────────────────────────────────────────────

class FormulaResponse(BaseModel):
    presupuesto_global: str = ""
    tiene_desglose: bool = False
    partidas: list = []
    resumen: str = ""
    agent: str = "agent-economico"

class SimulateRequest(BaseModel):
    partidas: list
    descuentos: list[str]
    num_empresas: int = 3
    session_id: str = "default"

class SimulateResponse(BaseModel):
    resultado: str
    agent: str = "agent-economico"

# ─── Text extraction helpers ─────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(p for p in parts if p.strip())
    except ImportError:
        raise HTTPException(status_code=422, detail="Librería pypdf no instalada.")
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Error leyendo PDF: {exc}")


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_texts = []
        for table in doc.tables:
            rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
            table_texts.append("\n".join(rows))
        return "\n\n".join(paragraphs + table_texts)
    except ImportError:
        raise HTTPException(status_code=422, detail="Librería python-docx no instalada.")
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Error leyendo DOCX: {exc}")


async def _extract_files(files: List[UploadFile]) -> tuple[str, str]:
    parts: list[str] = []
    names: list[str] = []
    for upload in files:
        file_name = upload.filename or ""
        ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
        if ext not in ("pdf", "docx"):
            raise HTTPException(status_code=415, detail=f"'{file_name}' no admitido. Solo .pdf y .docx.")
        data = await upload.read()
        if len(data) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail=f"'{file_name}' supera 50 MB.")
        text = _extract_pdf(data) if ext == "pdf" else _extract_docx(data)
        if text.strip():
            parts.append(f"=== DOCUMENTO: {file_name} ===\n\n{text}")
            names.append(file_name)
    if not parts:
        raise HTTPException(status_code=422, detail="No se pudo extraer texto de ningún fichero.")
    combined = "\n\n" + ("\n\n" + "─" * 60 + "\n\n").join(parts)
    return combined, " + ".join(names)

# ─── Endpoints ────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "agent": settings.agent_name, "version": settings.agent_version}


@app.post("/extract-formula", response_model=FormulaResponse)
async def extract_formula(
    files: List[UploadFile] = File(...),
    session_id: str = Form("default"),
):
    """Extrae la fórmula de valoración económica del PCAP."""
    text, names = await _extract_files(files)
    logging.info("extract-formula: %d chars, files=%s, session=%s", len(text), names, session_id)
    assert _agent is not None
    try:
        data = await _agent.extract_formula(
            document_text=text,
            session_id=session_id,
            file_name=names,
        )
    except Exception as exc:
        logging.error("Error en /extract-formula: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    return FormulaResponse(
        presupuesto_global=str(data.get("presupuesto_global", "")),
        tiene_desglose=bool(data.get("tiene_desglose", False)),
        partidas=data.get("partidas", []),
        resumen=str(data.get("resumen", "")),
    )


@app.post("/simulate", response_model=SimulateResponse)
async def simulate(req: SimulateRequest):
    """Simula escenarios económicos con los descuentos y partidas dados."""
    if not req.partidas:
        raise HTTPException(status_code=400, detail="Debes proporcionar al menos una partida económica.")
    assert _agent is not None
    try:
        resultado = await _agent.simulate(
            partidas=req.partidas,
            descuentos=req.descuentos,
            num_empresas=req.num_empresas,
            session_id=req.session_id,
        )
    except Exception as exc:
        logging.error("Error en /simulate: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    return SimulateResponse(resultado=resultado)
