"""
Agent Cualificacion — FastAPI entry point.
Specialized agent for opportunity qualification.

Endpoints:
  GET  /health                     — liveness check
  GET  /.well-known/agent-card.json — A2A agent discovery card
  POST /chat                        — conversational interface (multi-turn)
  POST /qualify                     — structured GO/NO-GO from uploaded documents
"""
import io
import logging
import uuid
from contextlib import asynccontextmanager
from typing import Annotated

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from shared.factory import get_auth_validator, get_llm, get_memory_store, get_settings

from .agent import CualificacionAgent

logger = logging.getLogger(__name__)
settings = get_settings()
_agent: CualificacionAgent | None = None


@asynccontextmanager
async def lifespan(app: FastAPI):
    global _agent
    _agent = CualificacionAgent(
        llm=get_llm(),
        memory=get_memory_store(),
        auth=get_auth_validator(),
    )
    logger.info("Agent Cualificacion started — model: %s", get_llm().model_id)
    yield


app = FastAPI(
    title="Akena — Agent Cualificacion",
    version=settings.agent_version,
    lifespan=lifespan,
)

app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])


# ── Schemas ───────────────────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    message: str
    session_id: str = ""
    stream: bool = False


class ChatResponse(BaseModel):
    session_id: str
    response: str


# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "agent": settings.agent_name, "version": settings.agent_version}


@app.get("/.well-known/agent-card.json")
async def agent_card():
    """A2A discovery: exposes agent skills to the orchestrator."""
    return {
        "name": "agent-cualificacion",
        "version": settings.agent_version,
        "type": "specialized",
        "description": "Evaluates and qualifies sales opportunities for Accenture Spain.",
        "skills": [
            {
                "id": "qualify_opportunity",
                "name": "Opportunity Qualification",
                "description": "Score and evaluate a sales opportunity against qualification criteria (BANT/MEDDIC).",
            },
            {
                "id": "generate_qualification_report",
                "name": "Qualification Report",
                "description": "Generate a structured GO/NO-GO qualification report from uploaded documents.",
            },
            {
                "id": "check_go_nogo",
                "name": "Go/No-Go Decision",
                "description": "Evaluate whether an opportunity meets the threshold to proceed.",
            },
        ],
    }


@app.post("/chat", response_model=ChatResponse)
async def chat(req: ChatRequest):
    """Conversational interface for qualification queries (multi-turn)."""
    session_id = req.session_id or str(uuid.uuid4())

    if req.stream:
        async def event_stream():
            async for chunk in _agent.stream(req.message, session_id):
                yield f"data: {chunk}\n\n"
        return StreamingResponse(event_stream(), media_type="text/event-stream")

    response = await _agent.run(req.message, session_id)
    return ChatResponse(session_id=session_id, response=response)


@app.post("/qualify")
async def qualify(
    files: Annotated[list[UploadFile], File(description="PDF or DOCX pliego documents")],
    doc_types: Annotated[list[str], Form(description="Type per file: administrativo | tecnico | anexo")],
    commercial_origin: Annotated[
        str,
        Form(
            description=(
                "Origen comercial (peso ~50% en GO/NO-GO): "
                "accenture_led | relationship_momentum | reactive_untracked"
            ),
        ),
    ] = "relationship_momentum",
):
    """
    Analyze uploaded procurement documents and return a structured GO/NO-GO result.

    Accepts multipart/form-data with:
    - files[]    — one or more PDF/DOCX files
    - doc_types[] — matching type per file (administrativo, tecnico, anexo)
    - commercial_origin — origen comercial declarado (mitad del criterio de decisión)

    Returns JSON matching the PrequalResult + extractedFields schema expected by the frontend.
    """
    if len(files) != len(doc_types):
        raise HTTPException(
            status_code=422,
            detail=f"files count ({len(files)}) must match doc_types count ({len(doc_types)})",
        )

    doc_texts: list[dict] = []
    for upload_file, doc_type in zip(files, doc_types):
        content = await upload_file.read()
        filename = upload_file.filename or "documento"
        text = _extract_text(filename, content)
        logger.info("qualify: extracted %d chars from '%s' (%s)", len(text), filename, doc_type)
        doc_texts.append({"name": filename, "doc_type": doc_type, "text": text})

    try:
        result = await _agent.qualify_documents(doc_texts, commercial_origin=commercial_origin)
    except Exception as exc:
        logger.error("qualify_documents failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"Analysis failed: {exc}") from exc

    return result


# ── Text extraction helpers ────────────────────────────────────────────────────

def _extract_text(filename: str, content: bytes) -> str:
    name_lower = filename.lower()
    if name_lower.endswith(".pdf"):
        return _extract_pdf(content)
    if name_lower.endswith(".docx"):
        return _extract_docx(content)
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return f"[No se pudo extraer texto de {filename}]"


def _extract_pdf(content: bytes) -> str:
    try:
        import pypdf  # type: ignore[import]
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        return f"[Error extrayendo PDF: {exc}]"


def _extract_docx(content: bytes) -> str:
    """Extrae párrafos y tablas (tarifas, perfiles, plazos suelen ir en tablas)."""
    try:
        import docx  # type: ignore[import]
        doc = docx.Document(io.BytesIO(content))
        parts: list[str] = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    parts.append(line)
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return f"[Error extrayendo DOCX: {exc}]"
