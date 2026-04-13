"""
FastAPI entry point — agent-pliego.
Endpoints:
  GET  /health
  GET  /.well-known/agent-card.json
  POST /chat            — conversational follow-up
  POST /analyze         — full pliego analysis from raw text
  POST /analyze-file    — full pliego analysis from PDF or DOCX upload
"""
import io
import logging
from contextlib import asynccontextmanager

from typing import List
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from shared.factory import get_llm, get_memory_store, get_settings
from src.agent import PligoAnalysisAgent

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

settings = get_settings()
_agent: PligoAnalysisAgent | None = None


@asynccontextmanager
async def lifespan(app: FastAPI):
    global _agent
    _agent = PligoAnalysisAgent(llm=get_llm(), memory=get_memory_store())
    yield


app = FastAPI(
    title="Akena — Agent Pliego",
    version=settings.agent_version,
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Request / Response models ────────────────────────────────────────────────

class ChatRequest(BaseModel):
    message: str
    session_id: str = "default"
    stream: bool = False

class ChatResponse(BaseModel):
    response: str
    agent: str = "agent-pliego"

class AnalyzeRequest(BaseModel):
    text: str
    depth: str = "medio"       # "breve" | "medio" | "extenso"
    session_id: str = "default"
    file_name: str = ""

class AnalyzeResponse(BaseModel):
    analysis: str
    depth: str
    agent: str = "agent-pliego"

# ─── Text extraction helpers ─────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(p for p in parts if p.strip())
    except ImportError:
        raise HTTPException(
            status_code=422,
            detail="Librería pypdf no instalada. Ejecuta: pip install pypdf",
        )
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Error leyendo el PDF: {exc}")


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_texts = []
        for table in doc.tables:
            rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
            table_texts.append("\n".join(rows))
        return "\n\n".join(paragraphs + table_texts)
    except ImportError:
        raise HTTPException(
            status_code=422,
            detail="Librería python-docx no instalada. Ejecuta: pip install python-docx",
        )
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Error leyendo el DOCX: {exc}")

# ─── Endpoints ────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "agent": settings.agent_name, "version": settings.agent_version}


@app.get("/.well-known/agent-card.json")
async def agent_card():
    return {
        "name": "agent-pliego",
        "version": settings.agent_version,
        "description": (
            "Analiza pliegos de contratación pública española: alcance, presupuesto, "
            "criterios de valoración, solvencia, perfiles y resumen ejecutivo."
        ),
        "capabilities": ["analyze", "analyze-file", "chat"],
        "endpoints": {
            "health":        f"http://localhost:{settings.port}/health",
            "chat":          f"http://localhost:{settings.port}/chat",
            "analyze":       f"http://localhost:{settings.port}/analyze",
            "analyze_file":  f"http://localhost:{settings.port}/analyze-file",
        },
    }


@app.post("/chat", response_model=ChatResponse)
async def chat(req: ChatRequest):
    """Continuación conversacional — el usuario hace preguntas sobre el pliego analizado."""
    if not req.message.strip():
        raise HTTPException(status_code=400, detail="El mensaje no puede estar vacío.")
    assert _agent is not None
    try:
        response = await _agent.run(req.message, req.session_id)
    except Exception as exc:
        logging.error("Error in /chat: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    return ChatResponse(response=response)


@app.post("/analyze", response_model=AnalyzeResponse)
async def analyze_text(req: AnalyzeRequest):
    """Analiza el texto del pliego enviado directamente en el body."""
    if not req.text.strip():
        raise HTTPException(
            status_code=400,
            detail="El campo 'text' está vacío. Proporciona el contenido del pliego.",
        )
    if len(req.text) > 500_000:
        raise HTTPException(status_code=413, detail="El texto supera el límite de 500 000 caracteres.")
    assert _agent is not None
    try:
        analysis = await _agent.analyze(
            document_text=req.text,
            depth=req.depth,
            session_id=req.session_id,
            file_name=req.file_name,
        )
    except Exception as exc:
        logging.error("Error in /analyze: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    return AnalyzeResponse(analysis=analysis, depth=req.depth)


@app.post("/analyze-file", response_model=AnalyzeResponse)
async def analyze_file(
    file: UploadFile = File(...),
    depth: str = Form("medio"),
    session_id: str = Form("default"),
):
    """Sube un fichero PDF o DOCX y recibe el análisis estructurado del pliego."""
    file_name = file.filename or ""
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

    if ext not in ("pdf", "docx"):
        raise HTTPException(status_code=415, detail="Solo se aceptan ficheros .pdf y .docx.")

    data = await file.read()
    if len(data) > 50 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="El fichero supera el límite de 50 MB.")

    document_text = _extract_pdf(data) if ext == "pdf" else _extract_docx(data)

    if not document_text.strip():
        raise HTTPException(
            status_code=422,
            detail=(
                "No se pudo extraer texto del fichero. "
                "Asegúrate de que no sea un PDF de imagen sin OCR."
            ),
        )

    logging.info(
        "analyze-file: '%s', %d chars, depth=%s, session=%s",
        file_name, len(document_text), depth, session_id,
    )
    assert _agent is not None
    try:
        analysis = await _agent.analyze(
            document_text=document_text,
            depth=depth,
            session_id=session_id,
            file_name=file_name,
        )
    except Exception as exc:
        logging.error("Error in /analyze-file: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))

    return AnalyzeResponse(analysis=analysis, depth=depth)


@app.post("/analyze-files", response_model=AnalyzeResponse)
async def analyze_files(
    files: List[UploadFile] = File(...),
    depth: str = Form("medio"),
    session_id: str = Form("default"),
):
    """Sube varios ficheros PDF/DOCX, extrae su texto y genera un único análisis combinado."""
    if not files:
        raise HTTPException(status_code=400, detail="Debes adjuntar al menos un fichero.")

    combined_parts: list[str] = []
    file_names: list[str] = []

    for upload in files:
        file_name = upload.filename or ""
        ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
        if ext not in ("pdf", "docx"):
            raise HTTPException(
                status_code=415,
                detail=f"Fichero '{file_name}' no admitido. Solo se aceptan .pdf y .docx.",
            )
        data = await upload.read()
        if len(data) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail=f"'{file_name}' supera el límite de 50 MB.")

        text = _extract_pdf(data) if ext == "pdf" else _extract_docx(data)
        if text.strip():
            combined_parts.append(f"=== DOCUMENTO: {file_name} ===\n\n{text}")
            file_names.append(file_name)

    if not combined_parts:
        raise HTTPException(
            status_code=422,
            detail="No se pudo extraer texto de ningún fichero. Comprueba que no sean PDFs de imagen sin OCR.",
        )

    document_text = "\n\n" + ("\n\n" + "─" * 60 + "\n\n").join(combined_parts)
    combined_name = " + ".join(file_names)

    logging.info(
        "analyze-files: %d ficheros (%s), %d chars, depth=%s, session=%s",
        len(file_names), combined_name, len(document_text), depth, session_id,
    )
    assert _agent is not None
    try:
        analysis = await _agent.analyze(
            document_text=document_text,
            depth=depth,
            session_id=session_id,
            file_name=combined_name,
        )
    except Exception as exc:
        logging.error("Error in /analyze-files: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc))

    return AnalyzeResponse(analysis=analysis, depth=depth)
